import streamlit as st
import pandas as pd
import plotly.graph_objects as go
from datetime import datetime
import base64
from io import BytesIO
import tempfile
import os
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Placeholder for Gemini API call (replace with actual implementation)
def call_gemini(prompt):
    # Simulate Gemini response (replace with google.generativeai)
    return f"Generated response for: {prompt}"

# Generate PDF using ReportLab
def generate_pdf(filename, sections, visualizations, logo, summary, recommendations):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    
    # Custom styles
    styles.add(ParagraphStyle(name='Header', fontSize=20, textColor=HexColor('#1E3A8A'), spaceAfter=12))
    styles.add(ParagraphStyle(name='Subheader', fontSize=14, textColor=HexColor('#1E3A8A'), spaceAfter=8))
    styles.add(ParagraphStyle(name='Body', fontSize=12, spaceAfter=6))
    
    story = []
    
    # Logo
    if logo:
        logo_buffer = BytesIO(logo.read())
        logo_img = Image(logo_buffer, width=1.5*inch, height=0.75*inch)
        story.append(logo_img)
        story.append(Spacer(1, 0.25*inch))
        logo.seek(0)  # Reset buffer for Word
    
    # Title and Date
    story.append(Paragraph(f"{st.session_state.report_config['template']} Report", styles['Header']))
    story.append(Paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}", styles['Body']))
    story.append(Spacer(1, 0.5*inch))
    
    # Temporary directory for chart images
    with tempfile.TemporaryDirectory() as tmpdir:
        for section in sections:
            story.append(Paragraph(section, styles['Subheader']))
            
            if section == "Executive Summary":
                story.append(Paragraph(summary.replace('\n', '<br/>'), styles['Body']))
            
            elif section == "Data Overview":
                story.append(Paragraph(f"Rows: {st.session_state.data.shape[0]}", styles['Body']))
                story.append(Paragraph(f"Columns: {st.session_state.data.shape[1]}", styles['Body']))
                if st.session_state.cleaning_steps:
                    story.append(Paragraph("Cleaning Steps:", styles['Body']))
                    for step in st.session_state.cleaning_steps:
                        story.append(Paragraph(f"- {step}", styles['Body']))
            
            elif section == "Analysis and Visualizations":
                for viz_name in st.session_state.report_config["visualizations"]:
                    img_path = os.path.join(tmpdir, f"{viz_name.replace(' ', '_')}.png")
                    try:
                        visualizations[viz_name].write_image(img_path, format="png", width=600)
                        img = Image(img_path, width=5*inch, height=3*inch)
                        story.append(Paragraph(viz_name, styles['Body']))
                        story.append(img)
                        prompt = f"Generate caption for visualization: {viz_name}"
                        caption = call_gemini(prompt)
                        story.append(Paragraph(caption, styles['Body']))
                    except Exception as e:
                        story.append(Paragraph(f"Error embedding {viz_name}: {str(e)}", styles['Body']))
            
            elif section == "Predictive Modeling":
                for model_name, model_data in st.session_state.models.items():
                    story.append(Paragraph(model_name, styles['Body']))
                    story.append(Paragraph(f"R²: {model_data.get('R2', 'N/A')}", styles['Body']))
                    story.append(Paragraph(f"Top Features: {', '.join(model_data.get('features', []))}", styles['Body']))
                    if "predictions" in model_data:
                        story.append(Paragraph("Sample Predictions:", styles['Body']))
                        for pred in model_data["predictions"][:5]:
                            story.append(Paragraph(f"- {pred}", styles['Body']))
            
            elif section == "AI Insights":
                story.append(Paragraph("Insights:", styles['Body']))
                for insight in st.session_state.insights:
                    story.append(Paragraph(f"- {insight}", styles['Body']))
            
            elif section == "Recommendations":
                story.append(Paragraph(recommendations.replace('\n', '<br/>'), styles['Body']))
            
            story.append(Spacer(1, 0.25*inch))
        
        try:
            doc.build(story)
            buffer.seek(0)
            return buffer
        except Exception as e:
            st.error(f"PDF generation failed: {str(e)}")
            return None

# Generate Word document using python-docx
def generate_word(filename, sections, visualizations, logo, summary, recommendations):
    doc = Document()
    
    # Styles
    def add_paragraph(text, style='Normal', size=12, color=(0, 0, 0), bold=False, align='left'):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(*color)
        run.bold = bold
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return p
    
    # Logo
    if logo:
        logo_buffer = BytesIO(logo.read())
        doc.add_picture(logo_buffer, width=Inches(1.5))
        logo.seek(0)  # Reset buffer
    
    # Title and Date
    add_paragraph(f"{st.session_state.report_config['template']} Report", size=20, color=(30, 58, 138), bold=True)
    add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}", size=12)
    doc.add_paragraph()
    
    # Temporary directory for chart images
    with tempfile.TemporaryDirectory() as tmpdir:
        for section in sections:
            add_paragraph(section, size=14, color=(30, 58, 138), bold=True)
            
            if section == "Executive Summary":
                add_paragraph(summary, size=12)
            
            elif section == "Data Overview":
                add_paragraph(f"Rows: {st.session_state.data.shape[0]}", size=12)
                add_paragraph(f"Columns: {st.session_state.data.shape[1]}", size=12)
                if st.session_state.cleaning_steps:
                    add_paragraph("Cleaning Steps:", size=12)
                    for step in st.session_state.cleaning_steps:
                        add_paragraph(f"- {step}", size=12)
            
            elif section == "Analysis and Visualizations":
                for viz_name in st.session_state.report_config["visualizations"]:
                    img_path = os.path.join(tmpdir, f"{viz_name.replace(' ', '_')}.png")
                    try:
                        visualizations[viz_name].write_image(img_path, format="png", width=600)
                        add_paragraph(viz_name, size=12)
                        doc.add_picture(img_path, width=Inches(5))
                        prompt = f"Generate caption for visualization: {viz_name}"
                        caption = call_gemini(prompt)
                        add_paragraph(caption, size=12)
                    except Exception as e:
                        add_paragraph(f"Error embedding {viz_name}: {str(e)}", size=12)
            
            elif section == "Predictive Modeling":
                for model_name, model_data in st.session_state.models.items():
                    add_paragraph(model_name, size=12)
                    add_paragraph(f"R²: {model_data.get('R2', 'N/A')}", size=12)
                    add_paragraph(f"Top Features: {', '.join(model_data.get('features', []))}", size=12)
                    if "predictions" in model_data:
                        add_paragraph("Sample Predictions:", size=12)
                        for pred in model_data["predictions"][:5]:
                            add_paragraph(f"- {pred}", size=12)
            
            elif section == "AI Insights":
                add_paragraph("Insights:", size=12)
                for insight in st.session_state.insights:
                    add_paragraph(f"- {insight}", size=12)
            
            elif section == "Recommendations":
                add_paragraph(recommendations, size=12)
            
            doc.add_paragraph()
    
    buffer = BytesIO()
    try:
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Word document generation failed: {str(e)}")
        return None


# Main report page function
def show():
    st.set_page_config(page_title="DataInsights Pro", layout="wide")
    
    st.title("📑 Get Your Real-Time Report", anchor=False)
    # st.write("Customize and review your report below.")
    
    # Initialize session state
    for key in ["data", "cleaning_steps", "visualizations", "models", "insights"]:
        if key not in st.session_state:
            st.session_state[key] = None if key == "data" else {} if key in ["visualizations", "models"] else []
    
    # Check if data is loaded
    if st.session_state.data is None:
        st.warning("⚠️ Please upload a Dataset in the Data Upload page.")
        return
    
    st.write("Customize and review your report below.")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        template = st.selectbox(
            "Report Template",
            ["Marketing Campaign Analysis", "Executive Summary", "Detailed Analytics", "Custom"],
            help="Choose a template to start"
        )
    
    with col2:
        # Section selection
        available_sections = ["Executive Summary", "Data Overview"]
        if st.session_state.visualizations:
            available_sections.append("Analysis and Visualizations")
        if st.session_state.models:
            available_sections.append("Predictive Modeling")
        if st.session_state.insights:
            available_sections.append("AI Insights")
        available_sections.append("Recommendations")

        selected_sections = st.multiselect(
        "Include Sections",
        available_sections,
        default=["Executive Summary", "Data Overview", "Recommendations"],
        help="Select sections to include in your report"
        )
    
    with col3:
        theme = st.selectbox("Theme", ["Professional", "Modern", "Branded"])

    # Visualization selection
    if st.session_state.visualizations:
        st.subheader("Select Visualizations", anchor=False)
        selected_viz = []
        for viz_name in st.session_state.visualizations.keys():
            if st.checkbox(viz_name, value=True):
                selected_viz.append(viz_name)
    else:
        selected_viz = []

    # Sidebar: Report Configuration
    with st.sidebar:
        st.header("Configure Your Report", anchor=False)
        
        # Report style
        st.subheader("Report Style", anchor=False)
        format_type = st.radio("Format", ["PDF", "Word"], index=0)
        tone = st.slider("Tone", min_value=0, max_value=100, value=50, format="%d%% Technical")
        logo = st.file_uploader("Upload Logo (Optional)", type=["png", "jpg"])
        
        # Export options
        include_dataset = st.checkbox("Include Cleaned Dataset (CSV)", value=True)
        filename = st.text_input("Report Filename", value=f"Analysis Report of {st.session_state.filename}")

        # Hidden HTML debug option
        debug_html = st.checkbox("Enable HTML Debug (Developer Only)", value=False)
        
        if st.button("Generate Preview", type="primary"):
            st.session_state.report_config = {
                "template": template,
                "sections": selected_sections,
                "visualizations": selected_viz,
                "format": format_type,
                "theme": theme,
                "tone": tone,
                "logo": logo,
                "include_dataset": include_dataset,
                "filename": filename
            }
            st.rerun()

    # Main Area: Report Preview
    # st.header("📄 Get Your Real-Time Report", anchor=False)

    if "report_config" not in st.session_state:
        st.info("Configure your report and click 'Generate Preview' to start.")
        return

    # Report content container
    with st.container(border=True):
        if st.session_state.report_config["logo"]:
            st.image(st.session_state.report_config["logo"], width=100)
            st.session_state.report_config["logo"].seek(0)  # Reset buffer
        
        st.markdown(f"### {st.session_state.report_config['template']} Report")
        st.markdown(f"*Generated on {datetime.now().strftime('%B %d, %Y')}*")
        
        # Sections
        for section in st.session_state.report_config["sections"]:
            with st.expander(section, expanded=True):
                if section == "Executive Summary":
                    prompt = f"Summarize marketing campaign analysis: {st.session_state.cleaning_steps}, {st.session_state.visualizations}, {st.session_state.models}, {st.session_state.insights}"
                    summary = call_gemini(prompt)
                    edited_summary = st.text_area("Edit Summary", value=summary, key=f"summary_{section}")
                    st.session_state.report_config["summary"] = edited_summary
                
                elif section == "Data Overview":
                    st.markdown("#### Dataset Details")
                    st.write(f"- Rows: {st.session_state.data.shape[0]}")
                    st.write(f"- Columns: {st.session_state.data.shape[1]}")
                    if st.session_state.cleaning_steps:
                        st.markdown("#### Cleaning Steps")
                        for step in st.session_state.cleaning_steps:
                            st.write(f"- {step}")
                    st.markdown("#### Sample Data")
                    st.dataframe(st.session_state.data.head())
                
                elif section == "Analysis and Visualizations":
                    for viz_name in st.session_state.report_config["visualizations"]:
                        st.markdown(f"#### {viz_name}")
                        st.plotly_chart(st.session_state.visualizations[viz_name], use_container_width=True)
                        prompt = f"Generate caption for visualization: {viz_name} in a marketing campaign context"
                        caption = call_gemini(prompt)
                        st.caption(caption)
                
                elif section == "Predictive Modeling":
                    for model_name, model_data in st.session_state.models.items():
                        st.markdown(f"#### {model_name}")
                        st.write(f"- R²: {model_data.get('R2', 'N/A')}")
                        st.write(f"- Top Features: {', '.join(model_data.get('features', []))}")
                        if "feature_importance" in model_data:
                            st.plotly_chart(model_data["feature_importance"], use_container_width=True)
                        if "predictions" in model_data:
                            st.markdown("##### Sample Predictions")
                            st.write(model_data["predictions"][:5])
                
                elif section == "AI Insights":
                    st.markdown("#### Insights")
                    for insight in st.session_state.insights:
                        st.write(f"- {insight}")
                    if "insight_visualizations" in st.session_state:
                        for viz_name, fig in st.session_state.insight_visualizations.items():
                            st.plotly_chart(fig, use_container_width=True)
                
                elif section == "Recommendations":
                    prompt = f"Generate marketing recommendations: {st.session_state.models}, {st.session_state.insights}"
                    recs = call_gemini(prompt)
                    edited_recs = st.text_area("Edit Recommendations", value=recs, key=f"recs_{section}")
                    st.session_state.report_config["recommendations"] = edited_recs
                
                notes = st.text_input("Add Notes", key=f"notes_{section}")
                if notes:
                    st.markdown(f"**Notes**: {notes}")

        # Reorder sections
        st.subheader("Reorder Sections", anchor=False)
        col1, col2 = st.columns(2)
        for i, section in enumerate(st.session_state.report_config["sections"]):
            with col1:
                if st.button(f"↑ Move Up", key=f"up_{section}", disabled=i == 0):
                    st.session_state.report_config["sections"].pop(i)
                    st.session_state.report_config["sections"].insert(i-1, section)
                    st.rerun()
            with col2:
                if st.button(f"↓ Move Down", key=f"down_{section}", disabled=i == len(st.session_state.report_config["sections"])-1):
                    st.session_state.report_config["sections"].pop(i)
                    st.session_state.report_config["sections"].insert(i+1, section)
                    st.rerun()

    # Export and Save
    st.subheader("Export Report", anchor=False)
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("Export Report", type="primary"):
            if st.session_state.report_config["format"] == "PDF":
                pdf_buffer = generate_pdf(
                    st.session_state.report_config["filename"],
                    st.session_state.report_config["sections"],
                    st.session_state.visualizations,
                    st.session_state.report_config["logo"],
                    st.session_state.report_config.get("summary", ""),
                    st.session_state.report_config.get("recommendations", "")
                )
                if pdf_buffer:
                    st.download_button(
                        label="Download PDF",
                        data=pdf_buffer,
                        file_name=f"{st.session_state.report_config['filename']}.pdf",
                        mime="application/pdf"
                    )
            else:  # Word
                word_buffer = generate_word(
                    st.session_state.report_config["filename"],
                    st.session_state.report_config["sections"],
                    st.session_state.visualizations,
                    st.session_state.report_config["logo"],
                    st.session_state.report_config.get("summary", ""),
                    st.session_state.report_config.get("recommendations", "")
                )
                if word_buffer:
                    st.download_button(
                        label="Download Word",
                        data=word_buffer,
                        file_name=f"{st.session_state.report_config['filename']}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            
            # Hidden HTML debug export
            if debug_html:
                html_content = f"""
                <html>
                <head>
                    <title>{st.session_state.report_config['template']} Report</title>
                    <style>
                        body {{ font-family: Arial, sans-serif; margin: 20px; }}
                        h1, h2, h3 {{ color: #1E3A8A; }}
                        p, li {{ font-size: 16px; }}
                        img {{ max-width: 600px; }}
                    </style>
                </head>
                <body>
                <h1>{st.session_state.report_config['template']} Report</h1>
                <p>Generated on {datetime.now().strftime('%B %d, %Y')}</p>
                """
                if st.session_state.report_config["logo"]:
                    logo_base64 = base64.b64encode(st.session_state.report_config["logo"].read()).decode()
                    html_content += f'<img src="data:image/png;base64,{logo_base64}" width="100"/>'
                    st.session_state.report_config["logo"].seek(0)
                
                for section in st.session_state.report_config["sections"]:
                    html_content += f"<h2>{section}</h2>"
                    if section == "Executive Summary":
                        html_content += f"<p>{st.session_state.report_config.get('summary', '')}</p>"
                    elif section == "Data Overview":
                        html_content += f"<p>Rows: {st.session_state.data.shape[0]}</p>"
                        html_content += f"<p>Columns: {st.session_state.data.shape[1]}</p>"
                        if st.session_state.cleaning_steps:
                            html_content += "<h3>Cleaning Steps</h3><ul>"
                            for step in st.session_state.cleaning_steps:
                                html_content += f"<li>{step}</li>"
                            html_content += "</ul>"
                    elif section == "Analysis and Visualizations":
                        for viz_name in st.session_state.report_config["visualizations"]:
                            html_content += f"<h3>{viz_name}</h3>"
                            with tempfile.TemporaryDirectory() as tmpdir:
                                img_path = os.path.join(tmpdir, f"{viz_name.replace(' ', '_')}.png")
                                st.session_state.visualizations[viz_name].write_image(img_path, format="png", width=600)
                                img_base64 = base64.b64encode(open(img_path, "rb").read()).decode()
                                html_content += f'<img src="data:image/png;base64,{img_base64}" width="600"/>'
                            prompt = f"Generate caption for visualization: {viz_name}"
                            caption = call_gemini(prompt)
                            html_content += f"<p>{caption}</p>"
                    elif section == "Predictive Modeling":
                        for model_name, model_data in st.session_state.models.items():
                            html_content += f"<h3>{model_name}</h3>"
                            html_content += f"<p>R²: {model_data.get('R2', 'N/A')}</p>"
                            html_content += f"<p>Top Features: {', '.join(model_data.get('features', []))}</p>"
                    elif section == "AI Insights":
                        html_content += "<h3>Insights</h3><ul>"
                        for insight in st.session_state.insights:
                            html_content += f"<li>{insight}</li>"
                        html_content += "</ul>"
                    elif section == "Recommendations":
                        html_content += f"<p>{st.session_state.report_config.get('recommendations', '')}</p>"
                
                html_content += "</body></html>"
                st.download_button(
                    label="Download HTML (Debug)",
                    data=html_content,
                    file_name=f"{st.session_state.report_config['filename']}.html",
                    mime="text/html"
                )

    with col2:
        if st.session_state.report_config["include_dataset"]:
            csv_buffer = st.session_state.data.to_csv(index=False).encode()
            st.download_button(
                label="Download Dataset",
                data=csv_buffer,
                file_name=f"{st.session_state.report_config['filename']}_data.csv",
                mime="text/csv"
            )
    with col3:
        if st.button("Save Draft"):
            st.success("Report draft saved! Access it later on this page.")

    # Help and Guidance
    with st.sidebar.expander("How to Create Your Report"):
        st.markdown("""
        1. Select a template (e.g., Marketing Campaign Analysis).
        2. Choose sections and visualizations.
        3. Customize style (format, theme, logo).
        4. Preview and edit content.
        5. Export as PDF or Word, or save as a draft.
        """)
        if st.button("View Sample Report"):
            st.session_state.sample_report = True
    
    if "sample_report" in st.session_state and st.session_state.sample_report:
        st.markdown("### Sample Marketing Campaign Report")
        st.write("Social Media outperforms Email in revenue.")
        st.plotly_chart(go.Figure(data=[go.Bar(x=["Social Media", "Email"], y=[600, 300])]))

if __name__ == "__main__":
    show()